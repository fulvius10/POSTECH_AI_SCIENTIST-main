"""Gera o documento Word autoexplicativo do desafio PayFlow."""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
FIGURES = REPORTS / "figures"
OUTPUT = ROOT / "docs" / "relatorio_payflow_crisp_dm.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(90, 98, 108)
LIGHT_GRAY = "F2F4F7"
RED = RGBColor(155, 28, 28)
GOLD = RGBColor(122, 90, 0)


def set_run_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_run_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    settings = [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]
    for name, size, color, before, after in settings:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left = p.add_run("POS TECH | CRISP-DM na Pratica")
    set_run_font(left, size=9, bold=True, color=DARK_BLUE)
    right = p.add_run("                                      PayFlow Credit Risk")
    set_run_font(right, size=9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])


def add_paragraph(doc, text, *, bold_prefix=None, italic=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True, color=color)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, italic=italic, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.keep_together = True
        run = p.add_run(item)
        set_run_font(run)


def add_callout(doc, label, text, color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=color)
    text_run = p.add_run(text)
    set_run_font(text_run)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths_dxa, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    repeat_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(header))
        set_run_font(run, size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, filename, caption, width=6.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline_shape = run.add_picture(str(FIGURES / filename), width=Inches(width))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("name", filename)
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_together = True
    r = cap.add_run(caption)
    set_run_font(r, size=9, italic=True, color=GRAY)


def page_break(doc):
    doc.add_page_break()


def main():
    metrics = json.loads((REPORTS / "metrics.json").read_text(encoding="utf-8"))
    quality = json.loads((REPORTS / "data_quality.json").read_text(encoding="utf-8"))
    comparison = pd.read_csv(REPORTS / "model_comparison.csv")
    importance = pd.read_csv(REPORTS / "feature_importance.csv").head(6)
    fairness_region = pd.read_csv(REPORTS / "fairness_by_region.csv")

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "PayFlow Credit Risk - CRISP-DM na pratica"
    doc.core_properties.subject = "Desafio da Disciplina - Score de Inadimplencia"
    doc.core_properties.author = "Projeto academico POS TECH"
    doc.core_properties.keywords = "CRISP-DM, credit risk, default, machine learning, FastAPI"

    # Capa: override nomeado editorial_cover_title.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(95)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(kicker.add_run("DESAFIO DA DISCIPLINA"), size=11, bold=True, color=GOLD)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("PayFlow Credit Risk"), size=30, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_run_font(subtitle.add_run("Do problema ao modelo: score de inadimplencia em 90 dias"), size=15, color=DARK_BLUE)
    add_paragraph(doc, "CRISP-DM na pratica | Modelo preditivo, API REST e estrategia de MLOps", italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=78)
    add_paragraph(doc, "POS TECH - AI Scientist | Fase 1", bold_prefix="POS TECH", color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Documento tecnico-executivo autoexplicativo", color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break(doc)
    doc.add_heading("Resumo executivo", level=1)
    add_callout(doc, "Recomendacao", "Usar o score para priorizar revisao humana; nao recusar credito automaticamente apenas pelo modelo.")
    add_paragraph(doc, f"A PayFlow precisa reduzir a inadimplencia inicial sem travar o crescimento. A base possui {quality['rows']:,} clientes, com {quality['default_count']} defaults ({quality['default_rate']:.2%}). O modelo estima a probabilidade de default em ate 90 dias usando somente informacoes disponiveis antes da concessao.")
    chosen = metrics["metrics_at_chosen_threshold"]
    summary_rows = [
        ("Modelo operacional", metrics["selected_model"]),
        ("ROC-AUC", f"{chosen['roc_auc']:.3f}"),
        ("Average Precision", f"{chosen['average_precision']:.3f} (prevalencia {metrics['holdout_default_rate']:.3f})"),
        ("Brier Score", f"{chosen['brier_score']:.3f} - menor e melhor"),
        ("Threshold didatico", f"{metrics['chosen_threshold']:.2f}"),
        ("Recall de default", f"{chosen['recall']:.1%}"),
    ]
    add_table(doc, ["Indicador", "Resultado"], summary_rows, [2700, 6660], font_size=10)
    add_paragraph(doc, "O threshold prioriza recall porque a premissa didatica considera um default nao identificado vinte vezes mais caro que uma revisao indevida. A decisao final depende de custos reais, politica de credito e governanca.")
    doc.add_heading("Fluxo CRISP-DM", level=2)
    add_numbered(doc, [
        "Entender a dor, o target, a decisao e as metricas de sucesso.",
        "Auditar qualidade, origem, disponibilidade temporal e representatividade dos dados.",
        "Preparar os dados sem leakage e dentro de pipelines reproduziveis.",
        "Comparar baseline e modelos candidatos, equilibrando performance e explicabilidade.",
        "Validar discriminacao, calibracao, erros, custo e equidade.",
        "Disponibilizar via API e operar com monitoramento, versionamento e retreinamento.",
    ])

    page_break(doc)
    doc.add_heading("1. Business Understanding", level=1)
    doc.add_heading("1.1 Dor, stakeholders e decisao", level=2)
    add_paragraph(doc, "A inadimplencia nos primeiros 90 dias gera perda financeira, pressiona a cobranca e deixa a diretoria receosa de expandir o credito. O modelo deve tornar a triagem mais consistente e escalavel, preservando revisao humana e regras de politica.")
    stakeholder_rows = [
        ("Diretoria e Risco", "Define apetite de risco e acompanha perda e crescimento."),
        ("Cobranca", "Planeja capacidade e prioriza clientes de maior risco."),
        ("Comercial e Produto", "Protege conversao e evita restricao indiscriminada."),
        ("Compliance/Juridico/DPO", "Valida LGPD, explicabilidade, equidade e contestacao."),
        ("Cliente", "Recebe decisao consistente, compreensivel e revisavel."),
    ]
    add_table(doc, ["Stakeholder", "Papel e dor"], stakeholder_rows, [2500, 6860])
    doc.add_heading("1.2 Definicao do target", level=2)
    add_bullets(doc, [
        "Unidade: cliente/proposta no momento da decisao de credito.",
        "Acao: default_90d = 1 quando o cliente entra em default em ate 90 dias; caso contrario, 0.",
        "Periodo: 90 dias, alinhado ao horizonte da dor e a maturacao do indicador.",
        "Momento da inferencia: antes da concessao. Dados posteriores sao proibidos.",
    ])
    doc.add_heading("1.3 KPIs de negocio x metas analiticas", level=2)
    kpi_rows = [
        ("Perda esperada 90d", "Reduzir EL = PD x LGD x EAD", "Calibracao, Brier, custo por threshold"),
        ("Defaults identificados", "Aumentar cobertura preventiva", "Recall e matriz de confusao"),
        ("Crescimento saudavel", "Manter aprovacao e margem", "Precision, taxa de revisao"),
        ("Carga operacional", "Controlar fila de analise/cobranca", "Falsos positivos e latencia"),
    ]
    add_table(doc, ["KPI", "Meta de negocio", "Suporte analitico"], kpi_rows, [2150, 3600, 3610], font_size=8.6)

    page_break(doc)
    doc.add_heading("2. Data Understanding", level=1)
    add_paragraph(doc, f"Foram analisadas {quality['rows']:,} observacoes e {quality['columns']} colunas. Nao ha linhas nem IDs duplicados. O target e desbalanceado: {quality['default_rate']:.2%} dos clientes entraram em default.")
    add_figure(doc, "01_target_distribution.png", "Figura 1 - Distribuicao do target default_90d.", width=5.7)
    doc.add_heading("2.1 Qualidade e ausencias", level=2)
    add_bullets(doc, [
        f"renda_mensal: {quality['missing_by_column']['renda_mensal']} ausencias ({quality['missing_by_column']['renda_mensal']/quality['rows']:.2%}).",
        f"tempo_emprego_anos: {quality['missing_by_column']['tempo_emprego_anos']} ausencias ({quality['missing_by_column']['tempo_emprego_anos']/quality['rows']:.2%}).",
        "Ausencia pode carregar informacao: autonomos ou clientes sem historico formal podem nao declarar tempo de emprego.",
    ])
    add_figure(doc, "02_missing_values.png", "Figura 2 - Colunas com valores ausentes.", width=5.4)
    doc.add_heading("2.2 Como os dados existiriam no mundo real", level=2)
    source_rows = [
        ("CRM/cadastro", "idade, renda declarada, emprego, regiao"),
        ("Core e motor de credito", "valor, prazo, juros, produto, decisao"),
        ("Bureau", "score, contratos, cartoes, atrasos anteriores"),
        ("App/site/parceiros", "canal de aquisicao e jornada"),
        ("Cobranca", "status e default observados apos 90 dias"),
    ]
    add_table(doc, ["Fonte", "Exemplos"], source_rows, [2700, 6660])

    page_break(doc)
    doc.add_heading("3. Data Preparation", level=1)
    doc.add_heading("3.1 Leakage: informacao do futuro", level=2)
    add_callout(doc, "Regra temporal", "Uma feature so pode entrar no modelo se existir no instante em que a proposta e analisada.", color=RED)
    leakage_rows = [
        ("parcelas_pagas_ate_3m", "So e conhecida tres meses apos a concessao."),
        ("atraso_primeira_parcela_dias", "Depende do pagamento posterior."),
        ("status_apos_90d", "Descreve diretamente o desfecho futuro."),
    ]
    add_table(doc, ["Coluna excluida", "Justificativa"], leakage_rows, [3500, 5860])
    leak = metrics["leakage_demo_metrics"]
    add_paragraph(doc, f"Como prova, um modelo treinado propositalmente com leakage obteve ROC-AUC {leak['roc_auc']:.1f}, Average Precision {leak['average_precision']:.1f} e acuracia {leak['accuracy']:.1f}. O resultado perfeito e um alerta de invalidez, nao uma conquista.")
    doc.add_heading("3.2 Tratamentos reproduziveis", level=2)
    add_bullets(doc, [
        "id_cliente removido: identificador sem significado preditivo.",
        "idade excluida do modelo e preservada somente para auditoria de equidade.",
        "Numericas: mediana do treino, indicador de ausencia e padronizacao.",
        "Categoricas: moda e one-hot encoding, aceitando categorias desconhecidas.",
        "Todos os passos ficam dentro de Pipeline/ColumnTransformer para impedir contaminacao do teste.",
    ])
    doc.add_heading("3.3 Outliers", level=2)
    add_paragraph(doc, "O criterio IQR foi usado como triagem. Nenhuma linha foi removida automaticamente: renda, valor solicitado e atrasos extremos podem representar riscos reais. Variaveis binarias e contagens infladas em zero exigem interpretacao de dominio, pois o IQR pode chamar valores validos de outliers.")

    page_break(doc)
    doc.add_heading("4. Modeling", level=1)
    add_paragraph(doc, "O desbalanceamento foi tratado pela escolha de metricas adequadas e pelo threshold, sem alterar artificialmente as probabilidades por class_weight. Foram comparados tres niveis de complexidade.")
    model_rows = []
    for _, row in comparison.iterrows():
        model_rows.append((row["model"], f"{row['cv_roc_auc_mean']:.3f}", f"{row['cv_average_precision_mean']:.3f}", f"{row['cv_brier_mean']:.3f}"))
    add_table(doc, ["Modelo", "CV ROC-AUC", "CV AP", "CV Brier"], model_rows, [3300, 2020, 2020, 2020], font_size=8.8)
    add_bullets(doc, [
        "Dummy: baseline que reproduz a prevalencia e testa se ha ganho real.",
        "Regressao logistica: interpretavel, linear e adequada como baseline de producao.",
        "Random Forest: captura nao linearidades e interacoes, com explicacao por permutacao.",
    ])
    add_callout(doc, "Escolha", f"{metrics['selected_base_model']} teve a maior Average Precision media. Em seguida, recebeu calibracao sigmoid usando apenas dados de treino.")
    add_figure(doc, "03_model_curves.png", "Figura 3 - Curvas ROC e Precisao-Recall dos modelos no holdout.", width=6.2)

    page_break(doc)
    doc.add_heading("5. Evaluation", level=1)
    doc.add_heading("5.1 Estrategia de validacao", level=2)
    add_paragraph(doc, "Foi utilizado holdout estratificado de 25% (1.250 clientes) e validacao cruzada estratificada de cinco partes no treino. O CSV nao possui data de concessao; logo, nao e possivel executar um split temporal verdadeiro. Em producao, o teste deve sempre ocorrer em meses posteriores aos usados no treino.")
    default = metrics["metrics_at_0_50"]
    chosen = metrics["metrics_at_chosen_threshold"]
    threshold_rows = [
        ("0,50", f"{default['precision']:.1%}", f"{default['recall']:.1%}", str(default["fp"]), str(default["fn"])),
        (f"{metrics['chosen_threshold']:.2f}", f"{chosen['precision']:.1%}", f"{chosen['recall']:.1%}", str(chosen["fp"]), str(chosen["fn"])),
    ]
    add_table(doc, ["Threshold", "Precision", "Recall", "FP", "FN"], threshold_rows, [1900, 1900, 1900, 1830, 1830], font_size=9)
    add_paragraph(doc, f"AUC {chosen['roc_auc']:.3f} mede ordenacao; AP {chosen['average_precision']:.3f} enfatiza a classe rara; Brier {chosen['brier_score']:.3f} mede a qualidade probabilistica. Nenhuma metrica isolada define sucesso de negocio.")
    add_figure(doc, "04_confusion_matrix.png", "Figura 4 - Matriz de confusao no threshold escolhido.", width=4.9)
    doc.add_heading("5.2 Calibracao", level=2)
    add_paragraph(doc, "Calibracao compara probabilidade prevista com frequencia observada. Ela e indispensavel quando o score sera interpretado como PD e usado em perda esperada. O Brier melhorou e a curva abaixo permite verificar desvios por faixa.")
    add_figure(doc, "07_calibration_curve.png", "Figura 5 - Curva de calibracao do modelo selecionado.", width=4.9)

    page_break(doc)
    doc.add_heading("6. Threshold e impacto de negocio", level=1)
    add_paragraph(doc, "O threshold nao deve ser escolhido por acuracia. Ele representa o ponto em que a organizacao transforma probabilidade em acao e, por isso, depende do custo dos erros.")
    add_callout(doc, "Premissa didatica", "Falso negativo = R$ 10.000; falso positivo = R$ 500. Substituir por EAD, LGD, margem e custo operacional reais.", color=GOLD)
    add_figure(doc, "05_threshold_cost.png", "Figura 6 - Custo ilustrativo calculado em previsoes out-of-fold do treino.", width=5.8)
    add_paragraph(doc, f"O minimo ocorreu em {metrics['chosen_threshold']:.2f}. No holdout, esse ponto capturou {chosen['tp']} dos {chosen['tp'] + chosen['fn']} defaults e deixou {chosen['fn']} passar, mas encaminhou {chosen['fp']} clientes sem default para revisao. A capacidade operacional precisa validar se essa fila e sustentavel.")
    doc.add_heading("Politica de decisao sugerida", level=2)
    add_bullets(doc, [
        "Abaixo do threshold: seguir politica de credito, sem aprovacao automatica garantida.",
        "Acima do threshold: revisao manual obrigatoria e analise de capacidade de pagamento.",
        "Probabilidades muito altas: segunda linha de validacao; jamais usar o score como justificativa unica.",
        "Registrar motivo, versao, decisao humana e eventual override para auditoria.",
    ])

    page_break(doc)
    doc.add_heading("7. Interpretacao e equidade", level=1)
    add_figure(doc, "06_feature_importance.png", "Figura 7 - Importancia por permutacao medida pela queda de Average Precision.", width=6.1)
    importance_rows = [(row["feature"], f"{row['importance_mean']:.3f}", f"{row['importance_std']:.3f}") for _, row in importance.iterrows()]
    add_table(doc, ["Variavel", "Importancia", "Desvio"], importance_rows, [5100, 2130, 2130], font_size=9)
    add_paragraph(doc, "Os sinais mais fortes estao ligados a historico de atraso, exposicao solicitada, renda e inadimplencias anteriores. Importancia nao implica causalidade e pode dividir credito entre variaveis correlacionadas.")
    doc.add_heading("7.1 Auditoria por grupos", level=2)
    fairness_rows = []
    for _, row in fairness_region.iterrows():
        fairness_rows.append((row["regiao"], int(row["records"]), f"{row['default_rate']:.1%}", f"{row['roc_auc']:.3f}", f"{row['mean_predicted_probability']:.3f}"))
    add_table(doc, ["Regiao", "N", "Default", "AUC", "PD media"], fairness_rows, [2450, 1200, 1700, 1700, 2310], font_size=8.7)
    add_paragraph(doc, "Diferencas regionais devem ser monitoradas, mas amostras menores aumentam incerteza. Idade foi excluida do modelo e usada apenas para auditoria por faixa. Antes de producao, devem ser avaliadas taxas de erro, calibracao e impacto adverso com apoio juridico.")

    page_break(doc)
    doc.add_heading("8. Deployment: API REST", level=1)
    add_paragraph(doc, "O artefato serializado contem preprocessamento, modelo calibrado, threshold, lista de features e versoes. Isso reduz divergencia entre treinamento e inferencia.")
    endpoint_rows = [
        ("GET /health", "Saude, model_version e target_version."),
        ("POST /predict", "PD, threshold, classe, faixa, acao e versoes."),
        ("/docs", "Swagger interativo gerado pelo FastAPI."),
    ]
    add_table(doc, ["Endpoint", "Contrato"], endpoint_rows, [2700, 6660])
    doc.add_heading("8.1 Fluxo de inferencia", level=2)
    add_numbered(doc, [
        "Canal envia os dados existentes antes da concessao.",
        "Pydantic valida tipos, limites, categorias e rejeita campos extras.",
        "Pipeline imputa, codifica e transforma com as estatisticas do treino.",
        "Modelo devolve a probabilidade calibrada de default em 90 dias.",
        "Threshold gera faixa e recomendacao de revisao humana.",
        "Servico registra versoes, score, decisao e latencia para monitoramento.",
    ])
    add_callout(doc, "Protecao contra leakage", "Se status_apos_90d ou outro campo extra for enviado, a API retorna HTTP 422.", color=RED)
    doc.add_heading("8.2 Execucao", level=2)
    add_paragraph(doc, "Treino: .\\.venv\\Scripts\\python.exe .\\src\\train_model.py", italic=True, color=DARK_BLUE)
    add_paragraph(doc, "API: .\\.venv\\Scripts\\python.exe -m uvicorn api.main:app --reload", italic=True, color=DARK_BLUE)
    add_paragraph(doc, "Teste: .\\.venv\\Scripts\\python.exe -m pytest -q", italic=True, color=DARK_BLUE)

    page_break(doc)
    doc.add_heading("9. Operacao e MLOps", level=1)
    mlops_rows = [
        ("Logging", "Request ID, timestamp, versoes, PD, faixa, decisao, override e latencia; minimizar PII."),
        ("Drift de entrada", "PSI/KS, categorias novas, ausentes e limites por feature."),
        ("Drift de score", "Distribuicao da PD e volume por faixa, produto e canal."),
        ("Performance", "Apos 90 dias: AUC, AP, Brier, recall, precision, custo e calibracao."),
        ("Equidade", "Erros e calibracao por grupos, com intervalos e revisao juridica."),
        ("Retreinamento", "Trimestral ou por gatilho de drift, deterioracao ou mudanca de politica."),
        ("Versionamento", "Dados, codigo, features, modelo, threshold e default_90d_v1."),
    ]
    add_table(doc, ["Pilar", "Implementacao proposta"], mlops_rows, [2450, 6910], font_size=9)
    doc.add_heading("9.1 Ciclo champion-challenger", level=2)
    add_numbered(doc, [
        "Treinar challenger com dados maduros e corte temporal.",
        "Validar tecnicamente, financeiramente, juridicamente e por equidade.",
        "Executar em shadow mode sem afetar decisoes.",
        "Realizar canary com pequeno percentual e criterios de rollback.",
        "Promover a champion somente com aprovacao formal e documentacao.",
    ])
    add_callout(doc, "Target versionado", "default_90d_v1 deve ter regra operacional, fonte, janela, data de vigencia e responsavel. Mudou a regra, muda a versao.")

    page_break(doc)
    doc.add_heading("10. Conclusao e limites", level=1)
    add_paragraph(doc, "O projeto cobre o fluxo solicitado: problema e stakeholders, metas, concepcao dos dados, ausentes, outliers, encoding, padronizacao, leakage, baseline, modelos candidatos, validacao, threshold de custo, API REST e MLOps.")
    doc.add_heading("O que o prototipo demonstra", level=2)
    add_bullets(doc, [
        "Separacao valida entre informacao presente e futura.",
        "Ganho real sobre o baseline para identificar clientes de risco.",
        "Probabilidade calibrada e threshold orientado a custo.",
        "Servico de inferencia reproduzivel e testado.",
        "Plano de monitoramento, retreinamento e governanca.",
    ])
    doc.add_heading("O que ainda impede uso real", level=2)
    add_bullets(doc, [
        "Base sintetica e ausencia de datas para validacao temporal.",
        "Custos financeiros ilustrativos, sem EAD/LGD/margem reais.",
        "Necessidade de validacao independente, juridica e de LGPD.",
        "Performance e equidade ainda nao observadas em producao.",
        "Ausencia de integracao com politica de credito, autenticacao e observabilidade corporativa.",
    ])
    add_callout(doc, "Conclusao", "A inteligencia esta no processo: um score util combina definicao correta do target, dados disponiveis no tempo certo, validacao honesta, decisao economica e operacao governada.")
    doc.add_heading("Artefatos entregues", level=2)
    add_bullets(doc, [
        "Codigo de treinamento e avaliacao em src/train_model.py.",
        "Modelo versionado em models/payflow_default_model.joblib.",
        "API FastAPI em api/main.py e testes em tests/test_api.py.",
        "Metricas, tabelas, previsoes e figuras em reports/ e data/processed/.",
        "Instrucoes de reproducao no README.md.",
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
